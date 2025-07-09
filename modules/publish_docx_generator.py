# modules/publish_docx_generator.py

"""Модуль для генерации DOCX-файлов со списками публикаций."""

from tkinter import filedialog, messagebox
from typing import Any, Dict

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd

from utils.docx_utils import replace_placeholders_in_para


def generate_docx(self: Any, name: str) -> None:
    """
    Генерирует DOCX-файл со списком публикаций на основе шаблона и данных.

    Args:
        self: Экземпляр главного окна приложения
        name: Название вкладки, для которой генерируется документ
    """
    try:
        # Получаем путь к шаблону
        template_path = getattr(self, f"{name}_template_path").get()
        if not template_path:
            messagebox.showerror("Ошибка", "Не выбран шаблон документа.")
            return

        # Создаем документ и получаем данные
        doc = Document(template_path)
        placeholders = self.placeholder_values.get(name, {})
        dataframe = self.dataframes.get(name)

        if dataframe is None or dataframe.empty:
            messagebox.showerror("Ошибка", "Нет данных для вставки.")
            return

        # Заменяем плейсхолдеры в документе
        for para in doc.paragraphs:
            replace_placeholders_in_para(para, placeholders)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        replace_placeholders_in_para(para, placeholders)

        # Вставляем список публикаций
        insert_list(doc, dataframe)

        # Сохраняем документ
        save_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Документы Word", "*.docx")],
            title="Сохранить список публикаций"
        )
        if save_path:
            doc.save(save_path)
            self.status.set(f"Файл сохранен: {save_path}")
    except Exception as e:
        messagebox.showerror(
            "Ошибка генерации",
            f"Произошла ошибка при создании документа:\n{str(e)}"
        )


def insert_list(doc: Document, dataframe: pd.DataFrame) -> None:
    """
    Вставляет список публикаций в документ на место маркера [[Список]].

    Args:
        doc: Объект документа для модификации
        dataframe: DataFrame с данными о публикациях
    """
    for para in doc.paragraphs:
        if '[[Список]]' in para.text:
            parent = para._element.getparent()
            index = parent.index(para._element)
            parent.remove(para._element)

            # Вставляем каждую публикацию как нумерованный пункт
            for i, (_, row) in enumerate(dataframe.iterrows(), start=1):
                new_para = doc.add_paragraph()
                new_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                new_para.paragraph_format.first_line_indent = Pt(18)  # Правильный отступ

                run = new_para.add_run(f"{i}. {row['Submitter']} {row['Title']}")
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

                parent.insert(index + i - 1, new_para._element)
            break

