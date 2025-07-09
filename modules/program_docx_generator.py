# modules/program_docx_generator.py

"""Модуль для генерации DOCX файла программы."""

from tkinter import filedialog, messagebox
from typing import Any

import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

from utils.docx_utils import replace_placeholders_in_para, months_ru


def generate_docx(self: Any, name: str) -> None:
    """
    Генерирует DOCX файл программы на основе шаблона и данных.

    Args:
        self: Экземпляр главного окна.
        name: Название вкладки, для которой генерируется документ.
    """
    try:
        template_path = getattr(self, f"{name}_template_path").get()
        if not template_path:
            messagebox.showerror("Ошибка", "Не выбран шаблон документа.")
            return

        doc = Document(template_path)
        placeholders = self.placeholder_values.get(name, {})
        dataframe = self.dataframes.get(name)

        if dataframe is None or dataframe.empty:
            messagebox.showerror("Ошибка", "Нет данных для вставки.")
            return

        # Замена плейсхолдеров в документе
        for para in doc.paragraphs:
            replace_placeholders_in_para(para, placeholders)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        replace_placeholders_in_para(para, placeholders)

        # Вставка списка докладов
        insert_list(doc, dataframe)

        # Сохранение документа
        save_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Documents", "*.docx")]
        )
        if save_path:
            doc.save(save_path)
            self.status.set(f"Файл сохранён: {save_path}")
    except Exception as e:
        messagebox.showerror(
            "Ошибка генерации",
            f"Произошла ошибка при генерации:\n{e}"
        )


def insert_list(doc: Document, dataframe: pd.DataFrame) -> None:
    """
    Вставляет список докладов в документ, группируя по датам.

    Args:
        doc: Объект документа docx.
        dataframe: DataFrame с данными о докладах.
    """
    # Преобразование и очистка данных
    dataframe['Дата и время начала'] = pd.to_datetime(
        dataframe['Дата и время начала'],
        errors='coerce'
    )
    dataframe = dataframe.dropna(subset=['Дата и время начала']).copy()
    dataframe['Дата'] = dataframe['Дата и время начала'].dt.date
    dataframe = dataframe.sort_values(['Дата', 'Дата и время начала'])
    grouped = dataframe.groupby('Дата')

    # Поиск места для вставки
    for para in doc.paragraphs:
        if '[[Список]]' in para.text:
            parent = para._element.getparent()
            index = parent.index(para._element)
            parent.remove(para._element)

            session_number = 1
            insert_pos = index

            # Вставка данных по сессиям
            for date, group_df in grouped:
                earliest_dt = group_df['Дата и время начала'].min()
                day = earliest_dt.day
                month = months_ru[earliest_dt.month]
                date_str = f"{day} {month}"
                time_str = earliest_dt.strftime("%H:%M")
                room = group_df.iloc[0]['Ауд.']

                # Заголовок сессии
                session_header = doc.add_paragraph()
                session_header.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run_header = session_header.add_run(f"\nЗаседание {session_number}.")
                run_header.font.name = 'Times New Roman'
                run_header.font.size = Pt(14)
                run_header.bold = True
                run_header._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                session_header.paragraph_format.space_after = Pt(2)

                # Информация о сессии
                session_info = doc.add_paragraph()
                session_info.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run_info = session_info.add_run(f"{date_str}, {time_str}, ауд. {room}\n")
                run_info.font.name = 'Times New Roman'
                run_info.font.size = Pt(12)
                run_info.bold = True
                run_info._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                session_info.paragraph_format.space_after = Pt(6)

                parent.insert(insert_pos, session_header._element)
                parent.insert(insert_pos + 1, session_info._element)
                insert_pos += 2

                # Вставка докладов
                for i, (_, row) in enumerate(group_df.iterrows(), start=1):
                    # Строка с докладчиком
                    p1 = doc.add_paragraph()
                    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run1 = p1.add_run(
                        f"\t{i}. {row['ФИО докладчика']}, группа {row['Номер группы']}"
                    )
                    run1.font.name = 'Times New Roman'
                    run1.font.size = Pt(14)
                    run1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                    p1.paragraph_format.space_after = Pt(2)

                    # Строка с названием доклада
                    p2 = doc.add_paragraph()
                    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run2 = p2.add_run(row['Название доклада'])
                    run2.font.name = 'Times New Roman'
                    run2.font.size = Pt(14)
                    run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                    p2.paragraph_format.space_after = Pt(6)

                    parent.insert(insert_pos, p1._element)
                    parent.insert(insert_pos + 1, p2._element)
                    insert_pos += 2

                session_number += 1
            break
