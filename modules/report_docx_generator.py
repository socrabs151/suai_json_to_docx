# modules/report_docx_generator.py

"""Модуль для генерации отчетов в формате DOCX."""

from tkinter import filedialog, messagebox
from typing import Any

import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

from utils.docx_utils import replace_placeholders_in_para, months_ru


def generate_docx(self: Any, name: str) -> None:
    """
    Генерирует отчет в формате DOCX на основе шаблона и данных.

    Args:
        self: Экземпляр главного окна приложения
        name: Название вкладки, для которой генерируется документ
    """
    try:
        # Получение пути к шаблону
        template_path = getattr(self, f"{name}_template_path").get()
        if not template_path:
            messagebox.showerror("Ошибка", "Не выбран шаблон документа.")
            return

        # Создание документа и обработка данных
        doc = Document(template_path)
        placeholders = self.placeholder_values.get(name, {})
        dataframe = self.dataframes.get(name)

        if dataframe is None or dataframe.empty:
            messagebox.showerror("Ошибка", "Нет данных для вставки.")
            return

        # Преобразование дат
        dataframe['Дата и время начала'] = pd.to_datetime(
            dataframe['Дата и время начала'],
            errors='coerce'
        )
        dataframe['Дата'] = dataframe['Дата и время начала'].dt.date
        dataframe['Время'] = dataframe['Дата и время начала'].dt.time

        # Замена плейсхолдеров
        for para in doc.paragraphs:
            replace_placeholders_in_para(para, placeholders)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        replace_placeholders_in_para(para, placeholders)

        # Вставка таблицы с данными
        insert_list(doc, dataframe)

        # Сохранение документа
        save_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Документы Word", "*.docx")],
            title="Сохранить отчет"
        )
        if save_path:
            doc.save(save_path)
            self.status.set(f"Файл сохранен: {save_path}")

    except Exception as e:
        messagebox.showerror(
            "Ошибка генерации",
            f"Произошла ошибка при создании отчета:\n{str(e)}"
        )


def set_table_borders(table: Any) -> None:
    """
    Устанавливает границы для таблицы.

    Args:
        table: Объект таблицы docx
    """
    tbl = table._element
    tbl_pr = tbl.tblPr
    borders_xml = f'''
    <w:tblBorders {nsdecls('w')}>
      <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>
      <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>
      <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>
      <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>
      <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>
      <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>
    </w:tblBorders>'''
    borders = parse_xml(borders_xml)
    tbl_pr.append(borders)


def set_column_width(cell: Any, width_inches: float) -> None:
    """
    Устанавливает ширину столбца таблицы.

    Args:
        cell: Ячейка таблицы
        width_inches: Ширина в дюймах
    """
    cell.width = Inches(width_inches)
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_w = parse_xml(
        f'<w:tcW {nsdecls("w")} w:w="{str(int(width_inches * 1440))}" w:type="dxa"/>'
    )
    tc_pr.append(tc_w)


def insert_list(doc: Document, dataframe: pd.DataFrame) -> None:
    """
    Вставляет список докладов в документ в виде таблицы.

    Args:
        doc: Объект документа docx
        dataframe: DataFrame с данными о докладах
    """
    # Подготовка данных
    dataframe = dataframe.dropna(subset=['Дата и время начала']).copy()
    dataframe = dataframe.sort_values(['Дата', 'Дата и время начала'])
    grouped = dataframe.groupby('Дата')

    # Поиск места для вставки
    for para in doc.paragraphs:
        if '[[Таблица]]' not in para.text:
            continue

        parent = para._element.getparent()
        index = parent.index(para._element)
        parent.remove(para._element)

        session_number = 1
        insert_pos = index

        # Обработка каждой группы (по датам)
        for date, group_df in grouped:
            earliest_dt = group_df['Дата и время начала'].min()
            day = earliest_dt.day
            month = months_ru[earliest_dt.month]
            year = earliest_dt.year
            date_str = f"{day} {month} {year} г."
            time_str = earliest_dt.strftime("%H:%M")
            room = group_df.iloc[0].get('Комната', '—')
            address = group_df.iloc[0].get('Адрес', 'ул. Б. Морская, д. 67')

            # Заголовок заседания
            p0 = doc.add_paragraph(f"\nЗаседание {session_number}")
            p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p0.runs[0].font.name = 'Times New Roman'
            p0.runs[0].font.size = Pt(14)
            p0.runs[0].bold = True
            p0.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

            # Информация о заседании
            p1 = doc.add_paragraph(f"{date_str}, {time_str}")
            p1.add_run(f"\n{address}, ауд. {room}")
            p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p1.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.bold = True
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

            # Подписи
            p2 = doc.add_paragraph("Научный руководитель секции – ")
            p3 = doc.add_paragraph("Секретарь – ")
            for p in [p2, p3]:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.runs[0].font.name = 'Times New Roman'
                p.runs[0].font.size = Pt(12)
                p.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

            # Заголовок таблицы
            p_title = doc.add_paragraph("\nСписок докладов\n")
            p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_title.runs[0].font.name = 'Times New Roman'
            p_title.runs[0].font.size = Pt(12)
            p_title.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

            # Создание таблицы
            table = doc.add_table(rows=1, cols=4)
            set_table_borders(table)

            # Настройка заголовков таблицы
            hdr_cells = table.rows[0].cells
            headers = ['№ п/п', 'ФИО докладчика и тема', 'Статус', 'Решение']
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                for para in hdr_cells[i].paragraphs:
                    for run in para.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
                        run.bold = True
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

            # Установка ширины столбцов
            col_widths = [0.4, 2.5, 2.5, 2.5]
            for i, cell in enumerate(hdr_cells):
                set_column_width(cell, col_widths[i])

            # Заполнение таблицы данными
            for i, (_, row) in enumerate(group_df.iterrows(), start=1):
                group_number = row.get('Номер группы', '')
                status = (
                    f"Магистрант гр. {group_number}" if group_number.endswith('М')
                    else f"Студент гр. {group_number}" if group_number else ''
                )

                row_cells = table.add_row().cells
                row_data = [
                    str(i),
                    f"{row['ФИО докладчика']}. {row['Название доклада']}",
                    status,
                    row.get('Решение', '')
                ]

                for j, (cell, data) in enumerate(zip(row_cells, row_data)):
                    cell.text = data
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(12)
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                    set_column_width(cell, col_widths[j])

            # Вставка элементов в документ
            elements = [p0, p1, p2, p3, p_title, table]
            for i, element in enumerate(elements):
                parent.insert(insert_pos + i, element._element)

            insert_pos += len(elements)
            session_number += 1

        break




