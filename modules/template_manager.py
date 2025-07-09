# modules/template_manager.py

"""Модуль для работы с шаблонами DOCX."""

from tkinter import filedialog, messagebox
from typing import Any, Set
import re
from docx import Document


def choose_template(self: Any, name: str) -> None:
    """
    Открывает диалог выбора шаблона DOCX и сохраняет путь к нему.

    Args:
        self: Экземпляр главного окна.
        name: Название вкладки, для которой выбирается шаблон.
    """
    path = filedialog.askopenfilename(filetypes=[("DOCX files", "*.docx")])
    if path:
        getattr(self, f"{name}_template_path").set(path)
        scan_template_for_placeholders(self, name, path)


def scan_template_for_placeholders(
    self: Any, name: str, template_path: str
) -> None:
    """
    Сканирует шаблон DOCX на наличие плейсхолдеров и извлекает их.

    Args:
        self: Экземпляр главного окна.
        name: Название вкладки.
        template_path: Путь к файлу шаблона.
    """
    try:
        doc = Document(template_path)
        placeholder_pattern = re.compile(r"\{([^{}]+?)\}")

        found_placeholders: Set[str] = set()

        # Найти все плейсхолдеры в документе
        for para in doc.paragraphs:
            found_placeholders.update(placeholder_pattern.findall(para.text))
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    found_placeholders.update(placeholder_pattern.findall(cell.text))

        self.placeholders[name] = set()
        self.placeholder_values[name] = {}

        for full_placeholder in found_placeholders:
            if "::" in full_placeholder:
                ph_name, default_value = full_placeholder.split("::", 1)
            else:
                ph_name, default_value = full_placeholder, ""

            ph_name = ph_name.strip()
            default_value = default_value.strip()

            self.placeholders[name].add(ph_name)

            # Используем значение из UI, если есть, иначе — стандартное
            tab_params = {
                "Название": getattr(self, f"{name}_event_name").get(),
            }
            self.placeholder_values[name][ph_name] = tab_params.get(
                ph_name, default_value
            )

        self.show_placeholders_in_tree(name)
    except Exception as e:
        messagebox.showerror("Ошибка", f"Не удалось обработать шаблон:\n{e}")
